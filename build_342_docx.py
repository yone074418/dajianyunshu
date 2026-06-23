from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("3.4.2_仿真功能实现整理.docx")
FIG_DIR = Path("figures")
FIG_DIR.mkdir(exist_ok=True)


figures = [
    ("3.30", 55, (134.7, 301.2, 460.5, 487.6), "运输货物介绍", "Transport cargo introduction", 4.9),
    ("3.31", 56, (95.0, 85.0, 510.0, 365.0), "运输组合方式", "Combined transport mode", 5.6),
    ("3.32", 56, (153.6, 541.3, 441.4, 724.5), "挂车参数选择图", "Trailer parameter selection diagram", 4.8),
    ("3.33", 57, (90.0, 192.0, 510.0, 270.0), "大件运输线路选择基本流程", "Large transport line selection basic process", 5.7),
    ("3.34", 57, (94.0, 376.0, 500.0, 695.5), "障碍点测量图", "Diagram of obstacle point measurement", 5.6),
    ("3.35", 58, (149.1, 357.1, 445.6, 519.2), "坡道通过性勘测图", "Slope passability survey diagram", 5.0),
    ("3.36", 58, (151.9, 562.8, 443.3, 723.2), "线路选择图", "Route selection diagram", 5.0),
    ("3.37", 59, (149.1, 337.0, 445.9, 471.6), "挂车液压支撑编点图", "Trailer hydraulic support system point location diagram", 5.3),
    ("3.38", 59, (158.1, 587.7, 436.5, 722.9), "阀门交互图", "Valve interaction diagram", 4.9),
    ("3.39", 60, (141.4, 237.0, 453.9, 407.1), "货物放置图", "Cargo placement plan", 5.0),
    ("3.40", 60, (139.9, 502.8, 455.4, 672.8), "选择绑扎器材图", "Select the lashing device diagram", 5.0),
    ("3.41", 61, (141.4, 177.0, 453.9, 347.1), "选择绑扎点图", "Select the binding point diagram", 5.0),
]


def crop_figures():
    doc = fitz.open("article.pdf")
    paths = {}
    for num, page_no, rect, *_ in figures:
        page = doc[page_no - 1]
        pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=fitz.Rect(rect), alpha=False)
        path = FIG_DIR / f"fig_{num.replace('.', '_')}.png"
        pix.save(path)
        paths[num] = path
    doc.close()
    return paths


def set_east_asian_font(run, font_name="宋体"):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_spacing(paragraph, after=6, before=0, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(0x2E, 0x74, 0xB5), 16, 8),
        ("Heading 2", 13, RGBColor(0x2E, 0x74, 0xB5), 12, 6),
        ("Heading 3", 12, RGBColor(0x1F, 0x4D, 0x78), 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("3.4.2 仿真功能实现整理")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    set_east_asian_font(run, "黑体")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("来源：《大件运输虚拟仿真实验教学系统设计与实现》3.4.2 节")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    set_east_asian_font(run)


def add_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    set_paragraph_spacing(p)
    run = p.add_run(text)
    set_east_asian_font(run)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    p.paragraph_format.keep_with_next = True
    return p


def add_figure(doc, fig_paths, num, zh, en, width):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(fig_paths[num]), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_together = True
    r1 = cap.add_run(f"图 {num}  {zh}\n")
    r2 = cap.add_run(f"Fig. {num}  {en}")
    for r in (r1, r2):
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        set_east_asian_font(r)


def add_footer(doc):
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build_docx(fig_paths):
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    add_heading(doc, "3.4.2 仿真功能实现")
    add_para(doc, "本实验将《浙江石化气化炉运输方案》导入系统中，以 Unity3D 为基础，运用虚拟仿真技术和三维建模技术，对运输相关模型及场景进行构造，将导入系统的运输方案编制成任务形式进行实验流程。本实验将实验流程分为运输任务及货物介绍、简单配车、路线勘测、车组确定、绑扎加固和货物运输六大步骤并分别进行仿真。下文进行仿真详解。")

    add_heading(doc, "（1）运输任务及货物介绍")
    add_para(doc, "系统首先进行运输货物选择，并介绍实验的运输任务，包括运输起始点、运输道路状况等内容，对后面车辆组合方式提供参考。然后介绍运输货物的尺寸、重量，鼠标点击可交互 360° 查看货物模型，如图 3.30 所示，对车辆选定提供参考。")
    add_figure(doc, fig_paths, "3.30", "运输货物介绍", "Transport cargo introduction", 4.9)

    add_heading(doc, "（2）简单配车")
    add_para(doc, "运输车组分为两类，一类是提供牵引力的牵引车，另一类是承载货物的挂车，两类车组通过组合方式形成整体实现货物的运输。系统将牵引车与挂车组合方式分为三类：带鹅颈的半挂牵引车、不带鹅颈的全挂牵引车、自行式液压轴线车。通过鼠标点击每类组合方式展示动画组合演示，并附有每类组合方式的优缺点，如图 3.31 所示，参考运输道路选择适当的组合方式。")
    add_figure(doc, fig_paths, "3.31", "运输组合方式", "Combined transport mode", 5.6)
    add_para(doc, "车辆组合方式确定后，要进行运输车组的选定。系统首先展示大件运输行业常用的两类牵引车，分别为 6x6 和 8x8 式，通过鼠标点击可查看三维模型及车辆尺寸、重量、动力性相关参数。然后展示挂车，同样可查看模型及参数。根据货物参数选择挂车的轴线数和纵列数，保证挂车能有足够的长宽度来承载货物，如图 3.32 所示。随后根据车货总重计算所需牵引力大小，选择牵引车的种类和数量，在满足运输性和经济性条件下初步配定实验的运输车组，最后展示所选车组的综合参数。")
    add_figure(doc, fig_paths, "3.32", "挂车参数选择图", "Trailer parameter selection diagram", 4.8)

    add_heading(doc, "（3）路线勘测")
    add_para(doc, "运输车组初步配定后，需要对运输路线进行勘测。为保证运输车组能顺利将货物送达目的地，需要对可选线路进行勘测，实地考察得到各条可选线路信息后确定路线上障碍点，通过校核运输车组通过性进行判定障碍点能否处置，最终结合经济性选择出最佳运输线路。其选择流程图如图 3.33 所示。")
    add_figure(doc, fig_paths, "3.33", "大件运输线路选择基本流程", "Large transport line selection basic process", 5.7)
    add_para(doc, "大件运输路线勘测中影响运输车组通行的常见障碍点包括高度、坡度、弯道等。系统对障碍点进行虚拟仿真，交互点击显示出三维模型及场景，随后引导使用测量工具对障碍点通过交互实现模拟测量，实现障碍点测量教学，如图 3.34 所示。")
    add_figure(doc, fig_paths, "3.34", "障碍点测量图", "Diagram of obstacle point measurement", 5.6)
    add_para(doc, "障碍点测量教学后，系统模拟出三条可选运输线路，点击各条线路进行逐一勘测。高度障碍点通过测量其高度，比较车货总高度进行判定是否通行；圆弧弯道障碍点测量转弯半径，系统将嵌入的圆弧弯道通行能力模型对选定运输车组长度等参数进行计算得出最小转弯半径，二者相比判定是否通行；直交弯道障碍点测量道路夹角、弯道入口宽度、出口宽度、内圆角半径，系统通过嵌入的直交弯道通行能力模型计算出最小弯道出口宽度，与测量的弯道出口宽度比较判定是否通行；坡度障碍点测量坡度最低点和最高点水平距离和垂直距离，从而计算出坡度值，系统通过嵌入的牵引力模型对选定车组动力性参数进行计算，得出可通行的最大坡度，与坡度进行比较判定是否通行，如图 3.35 所示。桥梁承载能力的判定较为复杂，需要设计院进行专业测量分析，故在系统中简化为给定桥梁承载能力，通过比较车货总重进行判定是否通行。通过三条线路虚拟勘测找出各自障碍点并进行测量判定是否通行，对不能通行处提出修改方案，若无法提出修改方案，则该条线路不能通行；最后根据三条线路勘测结果，选择出可通行线路作为实验运输线路，如图 3.36 所示。")
    add_figure(doc, fig_paths, "3.35", "坡道通过性勘测图", "Slope passability survey diagram", 5.0)
    add_figure(doc, fig_paths, "3.36", "线路选择图", "Route selection diagram", 5.0)

    add_heading(doc, "（4）车组确定")
    add_para(doc, "路线勘测结束后，系统返回到配车部分，根据选定线路上存在的障碍点通过交互操作进行车组重新配置。坡道通过性不足可以选择增加牵引车并重新校核车辆有效牵引力，确保能够通行。高度通过性有误可以选择改变车组悬架高度进行调整。")
    add_para(doc, "随后系统进入挂车拼接和液压支撑编点模块。在之前实验中根据货物尺寸选择合适的挂车轴线数和纵列数，在该模块需要将单轴线、单纵列挂车通过鼠标交互方式拼接成所选定的多纵列、多轴线挂车，拼接完成后需要对挂车的液压支撑点进行编制。挂车的液压支撑点对应液压回路的中心位置，通过编点，将挂车整体液压回路分为几处不同部分，从而保障车货总重分摊到挂车每一轴线上的荷载力小于挂车允许最大轴线载荷，使得车辆能够正常运输。首先系统采用三点编制的形式，给出挂车平面示意图并给出液压支撑编点坐标数据，通过交互操作标出三点在平面图的位置，并用不同颜色将三处液压回路框选出来，如图 3.37 所示。")
    add_figure(doc, fig_paths, "3.37", "挂车液压支撑编点图", "Trailer hydraulic support system point location diagram", 5.3)
    add_para(doc, "随后系统展示出挂车每轴线上液压回路阀门三维模型，通过交互操作选择旋转上下阀门将三部分液压回路彼此断开联通，从而确保液压支撑编点实现，如图 3.38 所示。")
    add_figure(doc, fig_paths, "3.38", "阀门交互图", "Valve interaction diagram", 4.9)
    add_para(doc, "最后通过轴线载荷模型计算出各液压回路的轴线载荷，与挂车最大轴线载荷进行比较，若超过最大轴线载荷则返回到挂车参数模块重新选择；若满足则正式确定车组。")

    add_heading(doc, "（5）货物装车")
    add_para(doc, "货物装车模块分为货物装车和货物绑扎加固。在进行货物装车操作时需要将货物放置在挂车上，通过交互操作使用键盘进行放置，确保货物重心对准车组中心位置，通过车尾的液压控制箱中液压表进行观测货物是否放置正确，如图 3.39 所示。")
    add_figure(doc, fig_paths, "3.39", "货物放置图", "Cargo placement plan", 5.0)
    add_para(doc, "货物放置完毕后，进行绑扎加固操作。系统给出本次实验所需的绑扎加固工具，包括橡胶垫、钢丝绳、紧固葫芦、安全带、梯子、防磨衬垫，如图 3.40 所示。")
    add_figure(doc, fig_paths, "3.40", "选择绑扎器材图", "Select the lashing device diagram", 5.0)
    add_para(doc, "通过交互操作按照绑扎加固顺序点击需要使用的工具拖向车辆，随即展示对应操作动画。首先需在货物下方铺置橡胶垫用来增大摩擦阻力；然后使用梯子、安全带到达货物顶端，并使用钢丝绳对货物进行绑扎，此处系统提供 4 处不同绑扎点供选择，选择合适的绑扎点位使钢丝绳在汽化炉两侧分别斜拉倒八字加固，并且与挂车板夹角不应大于 60°，如图 3.41 所示；再使用紧固葫芦对钢丝绳进行紧固；最后在货物与车辆接触处放置防磨衬垫以减少摩擦导致的车货损伤。")
    add_figure(doc, fig_paths, "3.41", "选择绑扎点图", "Select the binding point diagram", 5.0)

    add_heading(doc, "（6）货物运输")
    add_para(doc, "绑扎加固操作完成后，进入到货物运输模块。该模块主要是对路线选择进行复验，系统给定三条可选线路中仅有一条线路可修改后通行，如果选择不可通行线路，演示车组在障碍点不可通行的动画，进行重新选择线路；选择正确线路后进行动画演示车货运输，实验结束。")

    add_heading(doc, "整理小结")
    add_para(doc, "本节围绕大件运输虚拟仿真实验的核心操作流程，将运输任务、配车、路线勘测、车组确定、装车绑扎和运输复验串联为连续的实验任务。系统通过三维模型、交互测量、参数校核和动画反馈，把运输方案中的关键决策转化为可操作、可验证的虚拟仿真实验过程。")
    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    fig_paths = crop_figures()
    build_docx(fig_paths)
    print(OUT.resolve())
